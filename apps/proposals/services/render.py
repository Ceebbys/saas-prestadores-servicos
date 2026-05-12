"""Geração de saída da proposta: contexto compartilhado + PDF + DOCX.

Mantém um único ponto de montagem do contexto (`build_proposal_context`) reusado
por preview, PDF e DOCX, garantindo paridade visual e evitando duplicação.
"""
from __future__ import annotations

import io
import re
from typing import Optional

from django.template.loader import render_to_string
from django.utils import timezone

from apps.proposals.models import Proposal


def build_proposal_context(proposal: Proposal, request=None) -> dict:
    """Monta o contexto compartilhado entre preview/PDF/DOCX.

    Resolve a imagem de cabeçalho efetiva (proposta → template → empresa)
    e prepara campos exibíveis.
    """
    header_image = None
    if proposal.header_image:
        header_image = proposal.header_image
    elif (
        proposal.use_template_header_image
        and proposal.template_id
        and proposal.template.header_image
    ):
        header_image = proposal.template.header_image
    elif getattr(proposal.empresa, "logo", None):
        header_image = proposal.empresa.logo

    # RV05 #6/-F — rodapé: cascata proposta → template (footer_image agora existe no template)
    footer_image = None
    if proposal.footer_image:
        footer_image = proposal.footer_image
    elif proposal.template_id and proposal.template.footer_image:
        footer_image = proposal.template.footer_image

    items = list(proposal.items.all().order_by("order", "id"))
    # RV05 #5 — Múltiplas formas de pagamento. Mantém compat com payment_method legado.
    # RV05-H — usa .all() (paridade com template HTML): formas já vinculadas à
    # proposta têm direito histórico de aparecer no documento, mesmo se o admin
    # desativou a forma globalmente depois. Para remover definitivamente,
    # hard-delete a FormaPagamento.
    payment_methods = list(proposal.payment_methods.all())

    return {
        "proposal": proposal,
        "items": items,
        "header_image": header_image,
        "header_image_url": header_image.url if header_image else "",
        "footer_image": footer_image,
        "footer_image_url": footer_image.url if footer_image else "",
        "payment_methods": payment_methods,
        "lead": proposal.lead,
        "contato": getattr(proposal.lead, "contato", None) if proposal.lead else None,
        "empresa": proposal.empresa,
        "now": timezone.now(),
        "is_print": True,
        "request": request,
    }


def render_proposal_html(proposal: Proposal, request=None) -> str:
    """Renderiza o template print em HTML pronto para WeasyPrint ou preview."""
    ctx = build_proposal_context(proposal, request=request)
    return render_to_string("proposals/render/proposal_print.html", ctx, request=request)


def render_proposal_pdf(proposal: Proposal, request=None) -> bytes:
    """Gera bytes de PDF a partir do template print.

    Usa `apps.core.document_render.pdf.render_html_to_pdf` que aplica
    `media_url_fetcher` — resolve imagens `/media/*` direto do storage
    (não depende de HTTP roundtrip ou Caddy configurado).
    """
    from apps.core.document_render.pdf import render_html_to_pdf

    html = render_proposal_html(proposal, request=request)
    base_url = None
    if request is not None:
        base_url = request.build_absolute_uri("/")
    return render_html_to_pdf(html, base_url=base_url)


def _strip_html(text: Optional[str]) -> str:
    """Converte HTML rich em texto plano para uso em DOCX (perde formatação)."""
    if not text:
        return ""
    # remove tags
    plain = re.sub(r"<[^>]+>", "", text)
    # decodifica entidades simples
    plain = (
        plain.replace("&nbsp;", " ")
        .replace("&amp;", "&")
        .replace("&lt;", "<")
        .replace("&gt;", ">")
        .replace("&quot;", '"')
    )
    # normaliza whitespace
    plain = re.sub(r"\s+", " ", plain).strip()
    return plain


def render_proposal_docx(proposal: Proposal) -> bytes:
    """Gera bytes de DOCX estruturado.

    Limitação documentada: rich-text é renderizado como texto plano. Para
    layout fiel, usar PDF.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Cm

    ctx = build_proposal_context(proposal)
    doc = Document()

    # Margens
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # Cabeçalho com imagem (se houver path local)
    header_image = ctx.get("header_image")
    if header_image and hasattr(header_image, "path"):
        try:
            doc.add_picture(header_image.path, width=Cm(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            # Imagem inacessível ou formato não suportado por python-docx — ignora
            pass

    # Título e número
    title = doc.add_heading(proposal.title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Proposta {proposal.number}  •  {ctx['now']:%d/%m/%Y}")
    run.font.size = Pt(10)

    # Cliente
    contato = ctx.get("contato")
    lead = ctx.get("lead")
    if contato or lead:
        doc.add_heading("Cliente", level=2)
        nome = (contato.name if contato else lead.name) if (contato or lead) else ""
        if nome:
            doc.add_paragraph(f"Nome: {nome}")
        if contato:
            if contato.email:
                doc.add_paragraph(f"E-mail: {contato.email}")
            if contato.phone or contato.whatsapp:
                doc.add_paragraph(f"Telefone: {contato.phone or contato.whatsapp}")

    # Introdução
    intro = _strip_html(proposal.introduction)
    if intro:
        doc.add_heading("Introdução", level=2)
        doc.add_paragraph(intro)

    # Body
    body = _strip_html(proposal.body)
    if body:
        doc.add_heading("Conteúdo", level=2)
        doc.add_paragraph(body)

    # Itens
    items = ctx.get("items") or []
    if items:
        doc.add_heading("Itens", level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Descrição"
        hdr[1].text = "Qtd"
        hdr[2].text = "Unit."
        hdr[3].text = "Total"
        for it in items:
            row = table.add_row().cells
            row[0].text = it.description
            row[1].text = f"{it.quantity}"
            row[2].text = f"R$ {it.unit_price:.2f}".replace(".", ",")
            row[3].text = f"R$ {it.total:.2f}".replace(".", ",")

    # Totais
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Subtotal: R$ {proposal.subtotal:.2f}".replace(".", ","))
    run.font.size = Pt(11)
    if proposal.discount_percent:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f"Desconto: {proposal.discount_percent}%")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Total: R$ {proposal.total:.2f}".replace(".", ","))
    run.font.size = Pt(13)
    run.bold = True

    # RV05 #5 — Múltiplas formas de pagamento
    # RV05-H — .all() para paridade com template HTML (ver build_proposal_context)
    formas = list(proposal.payment_methods.all())
    if formas:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(
            "Formas de pagamento aceitas: "
            + " · ".join(f.nome for f in formas)
        )
    elif proposal.payment_method:
        # Fallback legado
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(proposal.get_payment_method_display() or "")

    if proposal.is_installment and proposal.installment_count:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f"Em até {proposal.installment_count}x.")

    # Termos
    terms = _strip_html(proposal.terms)
    if terms:
        doc.add_heading("Termos e Condições", level=2)
        doc.add_paragraph(terms)

    # Validade
    if proposal.valid_until:
        doc.add_paragraph()
        doc.add_paragraph(
            f"Válida até {proposal.valid_until:%d/%m/%Y}"
        ).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # RV05 #6 — Rodapé configurável (imagem + texto)
    footer_text = _strip_html(proposal.footer_content)
    has_footer = bool(footer_text) or (
        proposal.footer_image and hasattr(proposal.footer_image, "path")
    )
    if has_footer:
        doc.add_paragraph()  # separador
        if proposal.footer_image and hasattr(proposal.footer_image, "path"):
            try:
                doc.add_picture(proposal.footer_image.path, width=Cm(4))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass
        if footer_text:
            p = doc.add_paragraph(footer_text)
            for run in p.runs:
                run.font.size = Pt(9)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
