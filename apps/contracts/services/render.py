"""Render de contratos: preview HTML, PDF (WeasyPrint), DOCX (python-docx).

Reusa `apps.core.document_render.pdf.render_html_to_pdf` (mesmo url_fetcher
seguro usado em propostas). Mantém DOCX idiossincrático (python-docx puro).
"""
from __future__ import annotations

import io
import re
from typing import Optional

from django.template.loader import render_to_string
from django.utils import timezone

from apps.contracts.models import Contract


def build_contract_context(contract: Contract, request=None) -> dict:
    """Monta o contexto compartilhado entre preview/PDF/DOCX.

    Resolve a imagem de cabeçalho/rodapé efetiva (contrato → template) e
    prepara campos exibíveis. Mantém dual-read com `content` legado: se
    `body` vazio e `content` preenchido, usa o legado como fallback.
    """
    # Header image cascade: contract → template
    header_image = None
    if contract.header_image:
        header_image = contract.header_image
    elif contract.template_id and contract.template.header_image:
        header_image = contract.template.header_image

    footer_image = contract.footer_image or None
    if not footer_image and contract.template_id and contract.template.footer_image:
        footer_image = contract.template.footer_image

    # Dual-read: prefere body (rich-text); cai em content legado
    body = contract.body or contract.content or ""

    return {
        "contract": contract,
        "header_image": header_image,
        "header_image_url": header_image.url if header_image else "",
        "footer_image": footer_image,
        "footer_image_url": footer_image.url if footer_image else "",
        "body": body,
        "lead": contract.lead,
        "contato": getattr(contract.lead, "contato", None) if contract.lead else None,
        "empresa": contract.empresa,
        "now": timezone.now(),
        "is_print": True,
        "request": request,
    }


def render_contract_html(contract: Contract, request=None) -> str:
    """Renderiza o HTML do contrato (preview ou base do PDF)."""
    ctx = build_contract_context(contract, request=request)
    return render_to_string("contracts/render/contract_print.html", ctx, request=request)


def render_contract_pdf(contract: Contract, request=None) -> bytes:
    """Gera bytes de PDF via WeasyPrint + url_fetcher seguro do core."""
    from apps.core.document_render.pdf import render_html_to_pdf

    html = render_contract_html(contract, request=request)
    base_url = request.build_absolute_uri("/") if request else None
    return render_html_to_pdf(html, base_url=base_url)


def _strip_html(text: Optional[str]) -> str:
    """Converte HTML rich em texto plano para DOCX (limitação documentada)."""
    if not text:
        return ""
    plain = re.sub(r"<[^>]+>", "", text)
    plain = (
        plain.replace("&nbsp;", " ")
        .replace("&amp;", "&")
        .replace("&lt;", "<")
        .replace("&gt;", ">")
        .replace("&quot;", '"')
    )
    plain = re.sub(r"\s+", " ", plain).strip()
    return plain


def render_contract_docx(contract: Contract) -> bytes:
    """Gera bytes de DOCX estruturado.

    Limitação: rich-text vira texto plano. Para layout fiel, usar PDF.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Cm

    ctx = build_contract_context(contract)
    doc = Document()

    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    # Cabeçalho com imagem
    header_image = ctx.get("header_image")
    if header_image and hasattr(header_image, "path"):
        try:
            doc.add_picture(header_image.path, width=Cm(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

    # Título + número
    title = doc.add_heading(contract.title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Contrato {contract.number}  •  {ctx['now']:%d/%m/%Y}")
    run.font.size = Pt(10)

    # Cabeçalho rich
    header_text = _strip_html(contract.header_content)
    if header_text:
        doc.add_paragraph(header_text)

    # Cliente
    contato = ctx.get("contato")
    lead = ctx.get("lead")
    if contato or lead:
        doc.add_heading("Cliente", level=2)
        nome = (contato.name if contato else lead.name) if (contato or lead) else ""
        if nome:
            doc.add_paragraph(f"Nome: {nome}")
        if contato:
            if contato.email:
                doc.add_paragraph(f"E-mail: {contato.email}")
            if contato.cpf_cnpj:
                doc.add_paragraph(f"Documento: {contato.cpf_cnpj}")

    # Introdução
    intro = _strip_html(contract.introduction)
    if intro:
        doc.add_heading("Introdução", level=2)
        doc.add_paragraph(intro)

    # Body (dual-read)
    body = _strip_html(ctx.get("body", ""))
    if body:
        doc.add_heading("Conteúdo", level=2)
        doc.add_paragraph(body)

    # Termos
    terms = _strip_html(contract.terms)
    if terms:
        doc.add_heading("Termos e Condições", level=2)
        doc.add_paragraph(terms)

    # Valor + vigência
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Valor: R$ {contract.value:.2f}".replace(".", ","))
    run.font.size = Pt(12)
    run.bold = True

    if contract.start_date:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        end_text = (
            f" até {contract.end_date:%d/%m/%Y}" if contract.end_date else " (sem data fim)"
        )
        p.add_run(f"Vigência: {contract.start_date:%d/%m/%Y}{end_text}")

    # Rodapé
    footer_text = _strip_html(contract.footer_content)
    has_footer = bool(footer_text) or (
        contract.footer_image and hasattr(contract.footer_image, "path")
    )
    if has_footer:
        doc.add_paragraph()
        if contract.footer_image and hasattr(contract.footer_image, "path"):
            try:
                doc.add_picture(contract.footer_image.path, width=Cm(4))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass
        if footer_text:
            p = doc.add_paragraph(footer_text)
            for run in p.runs:
                run.font.size = Pt(9)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
